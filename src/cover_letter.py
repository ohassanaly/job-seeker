import os
import httpx
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches

def fetch_job_offer(url: str) -> str:
    
    r = httpx.get(url)
    soup = BeautifulSoup(r.text, "html.parser")
    text=""
    for h4 in soup.find_all("h4"):
        title = h4.get_text(strip=True)
        if title in ["Descriptif du poste", "Profil recherché"]:
            content_parts = []
            for sibling in h4.find_next_siblings():
                if sibling.name in ["h3", "h4"]:
                    break
                content_parts.append(sibling.get_text(" ", strip=True))
            text = text + "\n".join(content_parts)
            
    return(text)

def tailor_letter(template_text: str, job_text: str, SYSTEM_PROMPT:str, client: OpenAI) -> str:
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"## COVER LETTER TEMPLATE\n{template_text}\n\n"
                    f"## JOB OFFER\n{job_text}"
                ),
            },
        ],
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()

def write_docx(tailored_text: str, output_path: str):
    """
    Preserve the template's formatting: replace paragraph content in-place,
    then append any remaining paragraphs from the new text.
    """
    doc = Document()
    new_paragraphs = [p for p in tailored_text.split("\n")]
 
    # Clear all existing paragraphs and rewrite with template style
    # Strategy: keep the first paragraph's style as the base style
    base_style = doc.paragraphs[0].style if doc.paragraphs else None
 
    # Remove all paragraphs
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
 
    # Re-add paragraphs using the tailored text
    for line in new_paragraphs:
        para = doc.add_paragraph(line)
        if base_style:
            para.style = base_style
        para.paragraph_format.space_after = Pt(0)
 
    doc.save(output_path)
    print(f"✅  Tailored cover letter saved → {output_path}")

if __name__ == "__main__":
    
    load_dotenv()

    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    with open("data/resources/cover_template.txt") as f:
      cover_template = f.read()
    
    with open("data/resources/SYSTEM_PROMPT.txt") as f:
      SYSTEM_PROMPT = f.read()

    url= "https://www.welcometothejungle.com/fr/companies/mirakl/jobs/data-scientist-nlp-genai-catalog_paris"
    
    job_text = fetch_job_offer(url)

    cover_letter = tailor_letter(cover_template, job_text, SYSTEM_PROMPT, client)

    write_docx(cover_letter, f"data/on_demand/{"_".join(url.split("/")[5:])+ ".docx"}")